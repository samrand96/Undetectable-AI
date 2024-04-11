import argparse
from docx import Document
import json
from openai import OpenAI
import requests
import datetime
import ollama
from deep_translator import GoogleTranslator
import urllib.parse



# Initialize system as a dictionary to hold various properties
system = {}

client = OpenAI(api_key='sk-')

def text_translate(text, src_lang, target_lang):
    try:
        translated = GoogleTranslator(source=src_lang, target=target_lang).translate(text=text)
        return translated
    except Exception as e:
        print(f"An error occurred during translation: {e}")
        return text

def mix_translate(text, src_lang, target_lang):
    """
    Translate the given text from src_lang to target_lang and back to src_lang using googletrans.
    """
    try:
        # print(f"Before: {text}")
        translated = GoogleTranslator(source=src_lang, target=target_lang).translate(text=text)
        translated_back = GoogleTranslator(source=target_lang, target=src_lang).translate(text=translated)
        # print(F"After: {translated_back}")
        return translated_back
    except Exception as e:
        print(f"An error occurred during translation: {e}")
        # Return the original text if translation fails
        return text

# Placeholder for initializing ollama 
def setup_ollama(model):
    # Now we're modifying the global system dictionary
    system['model'] = model
    system['humanizer'] = """You are HUMIZER and DECOHERE, an AI designed to humize and decohere and disguise AI generated content such that it passes as human in academic voice. You are to respond to the input with its modification. Do NOT answer questions or modify meanings. Do NOT present yourself in your answer."""
    system['grammer'] = "Only Fix grammer problem without chaning the words and places of the sentence"

def humanize_openai(text):
    try:
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            # prompt=text,
            messages=[
                {"role": "system", "content": system['humanizer']},
                {"role": "user", "content": text}
            ],
            temperature=2,
            # max_tokens=100,
            top_p=1.0,
            frequency_penalty=0.0,
            presence_penalty=0.0,
            stream=False
        )
        return response.choices[0].message.content
    except Exception as e:
        print(f"An error occurred during connection: {e}")
        return text

def rewriter(text):
    try:
        response = ollama.generate(
            model=system['model'],
            raw=False, # standard question
            prompt=system['prompt'].format(question=text)
        )
        rewrited = response['response']
        rewrited = humanize(rewrited)
        return rewrited
    except Exception as e:
        print(f"An error occurred during connection: {e}")
        return text

def humanize(text):
    try:

        response = ollama.generate(
            model=system['model'],
            raw=False, # standard question
            stream=False,
            system=system['humanizer'],
            prompt=text,
            options={
                'temperature': 2 # ramp up
            }
        )
        return response['response']
    except Exception as e:
        print(f"Error occured during the connection: {e}")
        return text


def grammer(text):

    response = ollama.generate(
        model=system['model'],
        raw=False, # standard question
        stream=False,
        system=system['grammer'],
        prompt=text,
        options={
            'temperature': 2 # ramp up
        }
    )
    # print(response)
    return response['response']
        
def getScore(input_text):

    url = 'https://writer.com/wp-admin/admin-ajax.php'

    headers = {
          "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10.15; rv:124.0) Gecko/20100101 Firefox/124.0",
          "Accept": "text/plain, */*; q=0.01",
          "Accept-Language": "en-US,en;q=0.5",
          "Accept-Encoding": "gzip, deflate",
          "Content-Type": "application/x-www-form-urlencoded; charset=UTF-8",
          "X-Requested-With": "XMLHttpRequest",
          "Content-Length": "1434",
          "Origin": "https://writer.com",
          "Connection": "keep-alive",
          "Referer": "https://writer.com/ai-content-detector/",
          "Cookie": "mutiny.user.token=b6826b65-d682-4452-a9ee-1e1f48ce6c56; _gcl_au=1.1.647978909.1712778806; _ga_JW9S1XHSVS=GS1.1.1712798817.4.0.1712798817.60.0.0; _ga=GA1.1.1927691333.1712778807; OptanonConsent=isIABGlobal=false&datestamp=Thu+Apr+11+2024+08%3A10%3A19+GMT%2B0800+(China+Standard+Time)&version=6.26.0&landingPath=https%3A%2F%2Fwriter.com%2Fai-content-detector%2F&groups=1%3A1%2C0_175949%3A1%2C0_175955%3A0%2C2%3A1%2C0_175947%3A1%2C3%3A1%2C0_175950%3A1%2C4%3A1%2C0_175962%3A1%2C0_175967%3A1%2C0_175960%3A0%2C0_180698%3A0%2C0_175959%3A0%2C0_180699%3A0%2C0_180707%3A0%2C0_180701%3A1%2C0_182460%3A1%2C0_175963%3A1%2C0_180703%3A1%2C0_180693%3A1%2C0_180697%3A1%2C0_180695%3A1%2C0_180700%3A1%2C0_175964%3A1%2C0_180702%3A1%2C0_180694%3A1%2C0_180706%3A1%2C8%3A0; _an_uid=-1; _gd_visitor=30974554-ac8c-44ec-818e-2125e74ad0be; cb_user_id=null; cb_group_id=null; cb_anonymous_id=%22702959cd-4df4-4077-8558-941c16747767%22; _clck=mgjf0z%7C2%7Cfku%7C0%7C1561; _clsk=1b17mjv%7C1712798817474%7C1%7C1%7Ci.clarity.ms%2Fcollect; __hstc=117297753.3003116ee0b0cca88c4cd3d76bd3a068.1712778841576.1712781986398.1712794252947.3; hubspotutk=3003116ee0b0cca88c4cd3d76bd3a068; __hssrc=1; tk_ai=LnbeyXUl5NySbNxz7u%2F1cHQz; _gd_session=8db9c4c9-601d-478a-810e-aec2388544cd; _uetsid=02991820f77411ee905ea14ca4441505; _uetvid=02991360f77411eea2baf72a2f8d5537",
          "Sec-Fetch-Dest": "empty",
          "Sec-Fetch-Mode": "cors",
          "Sec-Fetch-Site": "same-origin"
        }

    input = urllib.parse.quote(input_text)
    # Define the payload as a dictionary
    payload = {"action":"ai_content_detector_v2","inputs": input,"token":""}

    # Send the POST request with JSON payload and headers
    response = requests.post(url, data=payload, headers=headers)

    # Check the response status
    if response.status_code == 200:
        resp = json.loads(response.text)
        score = resp['score']
        label = resp['label']
        print(resp)
        if(label == 'Human-Generated'):
            return True
        else:
            return False
    else:
        print("POST request failed with status code:", response.status_code)
        return None



def isHuman(input_text):
    # Define the URL
    url = 'https://api.zerogpt.com/api/detect/detectText'

    # Define headers with Content-Type
    headers = {
        'Accept': 'application/json, text/plain, */*',
        'Accept-Encoding': 'gzip, deflate, br',
        'Accept-Language': 'en-US,en;q=0.8',
        'Content-Type': 'application/json',
        'Origin': 'https://www.zerogpt.com',
        'Referer': 'https://www.zerogpt.com/',
        'Sec-Ch-Ua': '"Not A(Brand";v="99", "Brave";v="121", "Chromium";v="121"',
        'Sec-Ch-Ua-Mobile': '?0',
        'Sec-Ch-Ua-Platform': '"Linux"',
        'Sec-Fetch-Dest': 'empty',
        'Sec-Fetch-Mode': 'cors',
        'Sec-Fetch-Site': 'same-site',
        'Sec-Gpc': '1',
        'User-Agent': 'Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/121.0.0.0 Safari/537.36'
    }
    # Define the payload as a dictionary
    payload = {'input_text': input_text}

    # Convert payload to JSON format
    json_payload = json.dumps(payload)

    # Send the POST request with JSON payload and headers
    response = requests.post(url, data=json_payload, headers=headers)

    # Check the response status
    if response.status_code == 200:
        resp = json.loads(response.text)
        return int(resp['data']['isHuman'])
    else:
        print("POST request failed with status code:", response.status_code)
        return None


# Process the document
def process_docx(file_path, output_file_path, target_lang, grammar, chatgpt):
    
    doc = Document(file_path)
    new_doc = Document()

    src_lang = 'english'


    for paragraph in doc.paragraphs:

        if paragraph.text.strip() == "":
            continue

        # Changed to use rewriter function which applies both generating and humanizing
        output_text = paragraph.text

        if(getScore(output_text)):
            new_doc.add_paragraph(output_text)
            continue

        while True:
            
            output_text = mix_translate(output_text, src_lang, target_lang)
            

            if(getScore(output_text)):
                break

            if(chatgpt):
                output_text = humanize_openai(output_text)
            else:    
                output_text = humanize(output_text)

            if(getScore(output_text)):
                break

            new_output_text = mix_translate(output_text, src_lang, target_lang)
                        
            if(getScore(output_text)):
                output_text = new_output_text
                break

        
        if(grammar):
            output_text = grammer(output_text)

        new_doc.add_paragraph(output_text)

    new_doc.save(output_file_path)
   
    print(f"Processed document saved as {output_file_path}")

if __name__ == "__main__":
    
    parser = argparse.ArgumentParser(description="Process DOCX files to generate and decohere responses.")
    parser.add_argument("file_path", type=str, help="Path to the original DOCX file")
    parser.add_argument("--output_file_path", type=str, default="", help="Path to save the processed DOCX file (optional)")
    parser.add_argument("--model", type=str, default="dolphin-mistral:latest", help="Model name for ollama (optional)")
    parser.add_argument("--target_lang", type=str, default="arabic", help="Language code for intermediate translation (default: ru for Russian) (optional)")
    parser.add_argument("--grammar", type=bool, default=False, help="Corrent Grammar (default: Off)")
    parser.add_argument("--chatgpt", type=bool, default=False, help="Use ChatGPT Insstead (optional)")
    args = parser.parse_args()

    if not args.output_file_path:
        current_datetime = datetime.datetime.now()
        timestamp = str(int(current_datetime.timestamp()))
        args.output_file_path = args.file_path.replace('.docx', "_"+timestamp+".docx")

    # Setup and process the document
    setup_ollama(args.model)
    
    process_docx(args.file_path, args.output_file_path, args.target_lang, args.grammar, args.chatgpt)
